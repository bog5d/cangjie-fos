"""
解析机构发来的尽调清单（Excel/Word/PDF/粘贴文字）→ 结构化需求项列表。

设计原则：代码负责格式解析（读文件 → 纯文本），AI 只负责语义提取（识别哪些是真正的需求项）。
两步拆开能显著提升准确率，避免 AI 搞混表格结构。
"""
from __future__ import annotations
import json
import logging
import os
from pathlib import Path

logger = logging.getLogger(__name__)


def parse_checklist(source: str, source_type: str) -> list[dict]:
    """
    解析尽调清单。

    source: 文件路径（str）或粘贴的文字内容（str）
    source_type: "excel" | "word" | "pdf" | "text"

    返回：[{"item_no": "1", "category": "基本情况", "requirement": "验资报告"}, ...]
    """
    raw_text = _extract_raw_text(source, source_type)
    return _llm_extract_items(raw_text)


def _extract_raw_text(source: str, source_type: str) -> str:
    """第一步：代码读文件，转成纯文字，不依赖 AI。"""
    if source_type == "text":
        return source

    path = Path(source)
    if source_type == "excel":
        return _read_excel(path)
    elif source_type == "word":
        return _read_word(path)
    elif source_type == "pdf":
        return _read_pdf(path)
    return source


def _read_excel(path: Path) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    ws = wb.active
    lines: list[str] = []
    for i, row in enumerate(ws.iter_rows(values_only=True)):
        cells = [str(c).strip() for c in row if c is not None and str(c).strip() not in ("", "None")]
        if cells:
            lines.append(f"行{i + 1}: " + " | ".join(cells))
    wb.close()
    return "\n".join(lines)


def _read_word(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    lines: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _read_pdf(path: Path) -> str:
    import pdfplumber
    texts: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages[:15]:
            t = page.extract_text() or ""
            if t.strip():
                texts.append(t)
    return "\n".join(texts)


_CHUNK_SIZE = 4000
_CHUNK_OVERLAP = 300


def _llm_extract_items(raw_text: str) -> list[dict]:
    """
    第二步：AI 从纯文字中提取结构化需求项。
    长文本自动分块（4000字符/块，300字符重叠），去重合并。
    使用 dd_llm_client 统一管理 provider 配置 + 3次重试。
    """
    chunks = _split_into_chunks(raw_text, _CHUNK_SIZE, _CHUNK_OVERLAP)
    all_items: list[dict] = []
    seen: set[str] = set()

    for chunk in chunks:
        chunk_items = _llm_extract_chunk(chunk)
        for item in chunk_items:
            # 以需求文字前60字符做去重 key（处理重叠区域）
            key = item["requirement"][:60].strip().lower()
            if key and key not in seen:
                seen.add(key)
                all_items.append(item)

    # 重新按顺序编号
    for i, item in enumerate(all_items):
        item["item_no"] = str(i + 1)

    return all_items


def _split_into_chunks(text: str, chunk_size: int, overlap: int) -> list[str]:
    """将长文本分割为有重叠的块列表。"""
    if len(text) <= chunk_size:
        return [text]
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunks.append(text[start:end])
        if end == len(text):
            break
        start = end - overlap
    return chunks


def _llm_extract_chunk(chunk_text: str) -> list[dict]:
    """对单个文本块调用 LLM 提取需求项（可被测试 monkeypatch）。使用 dd_llm_client 统一重试。"""
    from cangjie_fos.services.dd_llm_client import get_dd_llm_client, call_with_retry

    client = get_dd_llm_client()
    prompt = f"""以下是一份机构发来的尽调清单原始内容（可能包含表头、大类标题、说明行等噪音）：

{chunk_text}

请提取所有具体的资料需求项（忽略大类标题行、表头行、空行、说明行）。
以 JSON 数组格式返回，每项格式：
{{"item_no": "序号", "category": "所属大类名称（如基本情况/财务/法务）", "requirement": "具体资料需求描述"}}

只返回 JSON 数组，不要任何解释或 markdown 标记："""

    def _call():
        resp = client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=3000,
            temperature=0,
        )
        return resp.choices[0].message.content.strip()

    raw = call_with_retry(_call, max_retries=3)

    # 清理 markdown 代码块
    if raw.startswith("```"):
        parts = raw.split("```")
        raw = parts[1] if len(parts) > 1 else raw
        if raw.lower().startswith("json"):
            raw = raw[4:]

    try:
        items = json.loads(raw.strip())
    except json.JSONDecodeError as e:
        logger.error("LLM chunk 解析失败: %s\n原文: %s", e, raw[:300])
        return []

    return [
        {
            "item_no": str(item.get("item_no", i + 1)),
            "category": str(item.get("category", "")),
            "requirement": str(item.get("requirement", "")),
        }
        for i, item in enumerate(items)
        if isinstance(item, dict) and item.get("requirement")
    ]
