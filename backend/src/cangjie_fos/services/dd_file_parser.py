"""从 PDF/Word/Excel/txt 文件中提取文字内容（纯工具函数，无 IO 副作用）。"""
from __future__ import annotations
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".txt", ".md"}

# 全文抽取时 PDF 最多遍历的页数（安全帽，防扫描件/超大 PDF 把索引拖死）
_FULL_MAX_PAGES = 80


def extract_text(file_path: Path, max_chars: int = 800) -> tuple[str, bool]:
    """
    从文件提取文字内容。
    返回 (text, readable)。readable=False 表示无法读取（加密/图片 PDF/不支持格式）。
    """
    suffix = file_path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        return "", False
    try:
        if suffix == ".pdf":
            return _extract_pdf(file_path, max_chars), True
        elif suffix in (".docx", ".doc"):
            return _extract_docx(file_path, max_chars), True
        elif suffix in (".xlsx", ".xls"):
            return _extract_excel(file_path, max_chars), True
        elif suffix in (".txt", ".md"):
            text = file_path.read_text(encoding="utf-8", errors="ignore")
            return text[:max_chars], True
    except Exception as e:
        logger.warning("无法读取文件 %s: %s", file_path.name, e)
    return "", False


def extract_full_text(file_path: Path, max_chars: int = 6000) -> tuple[str, bool]:
    """提取文件「全文」内容（供精判节点逐条核对正文），上限 max_chars。

    与 extract_text 的区别：
      - extract_text 只读前 3 页 / 800 字，够生成 20 字摘要即可；
      - extract_full_text 读全部页 / 默认 6000 字，供 LLM 精判核对内容是否满足需求。

    返回 (text, readable)。readable=False 表示加密/图片型 PDF/不支持格式
    （此类文件正文读不出，精判会自动跳过，仍可靠文件名+摘要参与粗筛）。

    注：本函数只做「文字层」快速抽取（pdfplumber 等）。扫描件/图片型 PDF 读不出字、
    加密文件读不了——这些「死角」由 dd_content_extractor.extract_for_verify 在精判节点
    按需补齐（解密 + MarkItDown + 图片 OCR），见 v1.17.0。本函数保持轻量，不引入重依赖。
    """
    suffix = file_path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        return "", False
    try:
        if suffix == ".pdf":
            # 页数安全帽：扫描件/纯图片 PDF 每页抽不出字、永远到不了 max_chars，
            # 若不限页会把上千页全遍历一遍，索引时卡死。最多看 _FULL_MAX_PAGES 页。
            return _extract_pdf(file_path, max_chars, max_pages=_FULL_MAX_PAGES), True
        elif suffix in (".docx", ".doc"):
            return _extract_docx(file_path, max_chars), True
        elif suffix in (".xlsx", ".xls"):
            return _extract_excel(file_path, max_chars, max_rows=200), True
        elif suffix in (".txt", ".md"):
            text = file_path.read_text(encoding="utf-8", errors="ignore")
            return text[:max_chars], True
    except Exception as e:
        logger.warning("无法读取文件全文 %s: %s", file_path.name, e)
    return "", False


def _extract_pdf(path: Path, max_chars: int, max_pages: int | None = 3) -> str:
    import pdfplumber
    texts: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        pages = pdf.pages if max_pages is None else pdf.pages[:max_pages]
        for page in pages:
            t = page.extract_text() or ""
            texts.append(t)
            if sum(len(x) for x in texts) >= max_chars:
                break
    return " ".join(texts)[:max_chars]


def _extract_docx(path: Path, max_chars: int) -> str:
    from docx import Document
    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    # 也读表格内容
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return " ".join(parts)[:max_chars]


def _extract_excel(path: Path, max_chars: int, max_rows: int = 30) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    ws = wb.active
    rows: list[str] = []
    for i, row in enumerate(ws.iter_rows(values_only=True)):
        cells = [str(c).strip() for c in row if c is not None and str(c).strip() not in ("", "None")]
        if cells:
            rows.append(f"行{i + 1}: " + " | ".join(cells))
        if i >= max_rows:
            break
    wb.close()
    return "\n".join(rows)[:max_chars]
