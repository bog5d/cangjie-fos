"""
投后季报模板解析器（gk 模式 · scenario='post_investment'）。

与尽调清单的区别：
  - 尽调：机构发来「需求清单」（一条条要什么材料）→ 我们匹配材料
  - 投后：机构发来「季报模板」（带【】空格 + 编号字段 + 叙述段落）→ 我们从材料里
          穷尽数据填进每个空格，填不上的留白，人工审核

设计原则（沿用 dd_checklist_parser）：代码负责格式解析，结构化提取「待填项」。
季报模板结构是确定性的（【】占位符 / N、字段： / 章节标题 / 叙述要求），
所以用纯代码提取比 LLM 更可靠，也更好测。

输出 item 形状与 dd_match_items 完全兼容（item_no/category/requirement），
这样后续匹配、草稿、导出引擎可直接复用，只是 scenario 标记为 post_investment。
额外带 field_kind 元数据，供导出时重建模板。
"""
from __future__ import annotations

import logging
import re
from pathlib import Path

logger = logging.getLogger(__name__)

# 章节标题：一、二、…（中文数字 + 顿号）
_SECTION_RE = re.compile(r"^([一二三四五六七八九十]+)、(.+)$")
# 编号字段：1、xxx：  或  1. xxx:
_FIELD_RE = re.compile(r"^\s*(\d+)\s*[、.]\s*(.+?)[:：]")
# 内联空格 【】 / 【 】 / 【  】
_BLANK_RE = re.compile(r"【\s*】")
# 叙述字数要求：不少于N字 / N字左右 / N字
_NARRATIVE_RE = re.compile(r"(不少于|大约|约|)\s*(\d{2,4})\s*字")


def parse_report_template(source: str, source_type: str) -> list[dict]:
    """
    解析投后季报模板 → 待填项列表。

    source: 文件路径（str）或粘贴文字
    source_type: "word" | "pdf" | "text"

    返回：[{"item_no","category","requirement","field_kind"}, ...]
      field_kind: "blank"(内联【】) | "field"(编号字段) | "narrative"(叙述段落)
    """
    raw_text = _extract_raw_text(source, source_type)
    return extract_fillable_items(raw_text)


def _extract_raw_text(source: str, source_type: str) -> str:
    if source_type == "text":
        return source
    path = Path(source)
    if source_type == "word":
        return _read_word(path)
    if source_type == "pdf":
        return _read_pdf(path)
    return source


def _read_word(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    lines: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            line = " | ".join(cells)
            if line.strip(" |"):
                lines.append("[表]" + line)
    return "\n".join(lines)


def _read_pdf(path: Path) -> str:
    import pdfplumber
    out: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            out.append(page.extract_text() or "")
    return "\n".join(out)


def extract_fillable_items(text: str) -> list[dict]:
    """从季报模板纯文本里提取所有「待填项」。

    三类待填项：
      1. blank     —— 行内 【】 占位符（每个空格一项，带整句上下文 + 同句序号）
      2. field     —— 编号字段「N、xxx：」且冒号后为空（如「2、地址：」）
      3. narrative —— 叙述段落「N、xxx（不少于N字）」（需要写一段文字）
    章节标题（一、二、…）不入项，只作为后续项的 category。
    """
    items: list[dict] = []
    current_section = "未分类"
    seq = 0

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        # 章节标题 → 更新分组，不入项
        sec = _SECTION_RE.match(line)
        if sec:
            current_section = f"{sec.group(1)}、{sec.group(2).strip()}"
            # 章节标题本身若带叙述字数要求（如「三、核心能力分析（不少于500字）」）
            if _NARRATIVE_RE.search(line):
                seq += 1
                items.append({
                    "item_no": str(seq),
                    "category": current_section,
                    "requirement": _clean_requirement(line),
                    "field_kind": "narrative",
                })
            continue

        # 行内 【】 空格 → 每个空格一项，带整句上下文
        blanks = list(_BLANK_RE.finditer(line))
        if blanks:
            total = len(blanks)
            for idx, _ in enumerate(blanks, start=1):
                seq += 1
                ctx = _clean_requirement(line)
                req = ctx if total == 1 else f"{ctx}（第{idx}/{total}个空格）"
                items.append({
                    "item_no": str(seq),
                    "category": current_section,
                    "requirement": req,
                    "field_kind": "blank",
                })
            continue

        # 编号字段「N、xxx：」冒号后为空 → field（值非空说明是示例/已填，跳过）
        fld = _FIELD_RE.match(line)
        if fld:
            after_colon = re.split(r"[:：]", line, maxsplit=1)
            tail = after_colon[1].strip() if len(after_colon) > 1 else ""
            # 叙述型字段（带字数要求）
            if _NARRATIVE_RE.search(line):
                seq += 1
                items.append({
                    "item_no": str(seq),
                    "category": current_section,
                    "requirement": _clean_requirement(line),
                    "field_kind": "narrative",
                })
            elif tail in ("", "（）", "()"):
                seq += 1
                items.append({
                    "item_no": str(seq),
                    "category": current_section,
                    "requirement": fld.group(2).strip(),
                    "field_kind": "field",
                })
            # tail 非空（如「IPO中介：券商、会计师、律师」是说明）→ 仍作为 field 待填
            else:
                seq += 1
                items.append({
                    "item_no": str(seq),
                    "category": current_section,
                    "requirement": _clean_requirement(line),
                    "field_kind": "field",
                })
    return items


def _clean_requirement(line: str) -> str:
    """整理 requirement 文本：去掉表格前缀、压缩空白。"""
    line = line.replace("[表]", "").strip()
    line = re.sub(r"\s+", " ", line)
    return line
