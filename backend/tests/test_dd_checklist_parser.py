"""Tests for dd_checklist_parser — all LLM calls mocked."""
from __future__ import annotations
from unittest.mock import patch
import pytest


_MOCK_ITEMS = [
    {"item_no": "1", "category": "基本情况", "requirement": "实收资本验资报告"},
    {"item_no": "2", "category": "基本情况", "requirement": "营业执照"},
]


def _fake_llm_extract(raw_text: str) -> list[dict]:
    return _MOCK_ITEMS


def test_parse_text_input():
    from cangjie_fos.services.dd_checklist_parser import parse_checklist
    with patch("cangjie_fos.services.dd_checklist_parser._llm_extract_items", side_effect=_fake_llm_extract):
        result = parse_checklist("1. 验资报告\n2. 营业执照", "text")
    assert len(result) == 2
    assert result[0]["requirement"] == "实收资本验资报告"
    assert result[0]["item_no"] == "1"


def test_parse_excel_file(tmp_path):
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["序号", "内容"])
    ws.append(["1", "实收资本验资报告"])
    ws.append(["2", "营业执照"])
    path = tmp_path / "dd.xlsx"
    wb.save(str(path))

    from cangjie_fos.services.dd_checklist_parser import parse_checklist
    with patch("cangjie_fos.services.dd_checklist_parser._llm_extract_items", side_effect=_fake_llm_extract):
        result = parse_checklist(str(path), "excel")
    assert len(result) == 2


def test_parse_word_file(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_paragraph("一、基本情况")
    doc.add_paragraph("1. 实收资本验资报告")
    doc.add_paragraph("2. 营业执照")
    path = tmp_path / "dd.docx"
    doc.save(str(path))

    from cangjie_fos.services.dd_checklist_parser import parse_checklist
    with patch("cangjie_fos.services.dd_checklist_parser._llm_extract_items", side_effect=_fake_llm_extract):
        result = parse_checklist(str(path), "word")
    assert len(result) == 2


def test_items_have_required_fields():
    from cangjie_fos.services.dd_checklist_parser import parse_checklist
    with patch("cangjie_fos.services.dd_checklist_parser._llm_extract_items", side_effect=_fake_llm_extract):
        result = parse_checklist("任意文字", "text")
    for item in result:
        assert "item_no" in item
        assert "category" in item
        assert "requirement" in item
        assert item["requirement"]  # 非空


# ── dd_match_service 测试 ────────────────────────────────────

def test_create_match_session_stores_items():
    from cangjie_fos.services.dd_match_service import create_match_session, get_session_items
    items = [
        {"item_no": "1", "category": "基本情况", "requirement": "验资报告"},
        {"item_no": "2", "category": "财务", "requirement": "审计报告"},
    ]
    session_id = create_match_session("test_tenant", "测试清单.xlsx", "/some/folder", items)
    assert session_id

    stored = get_session_items(session_id)
    assert len(stored) == 2
    assert stored[0]["requirement"] == "验资报告"
    assert stored[1]["category"] == "财务"


def test_run_matching_updates_confidence():
    from cangjie_fos.services.dd_match_service import create_match_session, run_matching, get_session_items
    from unittest.mock import patch

    items = [{"item_no": "1", "category": "基本情况", "requirement": "验资报告"}]
    session_id = create_match_session("test_tenant", "test.xlsx", "/folder", items)

    fake_index = [{"file_path": "/folder/验资.pdf", "filename": "验资.pdf", "summary": "验资报告"}]

    def fake_batch_match(items, file_list_text, index_rows):
        return {items[0]["id"]: {"file_path": "/folder/验资.pdf", "filename": "验资.pdf",
                                  "confidence": 0.92, "reason": "文件名匹配"}}

    with patch("cangjie_fos.services.dd_match_service._llm_batch_match", side_effect=fake_batch_match):
        with patch("cangjie_fos.services.dd_match_service._get_index_for_folder", return_value=fake_index):
            run_matching(session_id, "/folder")

    result = get_session_items(session_id)
    assert result[0]["confidence"] == pytest.approx(0.92)
    assert result[0]["matched_filename"] == "验资.pdf"


# ── dd_export_service 测试 ────────────────────────────────────

def test_export_creates_folder_and_copies_files(tmp_path):
    from cangjie_fos.services.dd_match_service import create_match_session, get_session_items
    from cangjie_fos.services.dd_export_service import export_to_folder
    from cangjie_fos.services.db_base import _connect
    from pathlib import Path

    # 准备真实文件
    src = tmp_path / "验资报告.pdf"
    src.write_bytes(b"fake pdf")

    items = [
        {"item_no": "1", "category": "基本情况", "requirement": "验资报告"},
        {"item_no": "2", "category": "基本情况", "requirement": "军工四证"},
    ]
    session_id = create_match_session("test", "dd.xlsx", str(tmp_path), items)

    # 手动给 item 1 设置匹配结果，item 2 标记 user_skipped
    stored = get_session_items(session_id)
    item1_id, item2_id = stored[0]["id"], stored[1]["id"]
    with _connect() as conn:
        conn.execute(
            "UPDATE dd_match_items SET matched_file_path=?, matched_filename=?, confidence=? WHERE id=?",
            (str(src), "验资报告.pdf", 0.95, item1_id),
        )
        conn.execute("UPDATE dd_match_items SET user_skipped=1 WHERE id=?", (item2_id,))

    out_dir = str(tmp_path / "output")
    result = export_to_folder(session_id, out_dir)

    assert result["exported"] == 1
    assert result["missing"] == 1

    # 文件被复制进去
    copied = list(Path(out_dir).rglob("*验资报告.pdf"))
    assert len(copied) == 1

    # 缺失清单生成
    gap = Path(out_dir) / "缺失清单.txt"
    assert gap.exists()
    assert "军工四证" in gap.read_text(encoding="utf-8")


def test_long_checklist_splits_into_multiple_chunks(monkeypatch):
    """超过4000字的清单应分多块处理，两块都应有 LLM 调用。"""
    call_count = 0

    def mock_extract_chunk(chunk_text: str) -> list[dict]:
        nonlocal call_count
        call_count += 1
        # 每块返回一个不同的需求项
        return [{"item_no": str(call_count), "category": "测试", "requirement": f"需求第{call_count}块"}]

    monkeypatch.setattr(
        "cangjie_fos.services.dd_checklist_parser._llm_extract_chunk",
        mock_extract_chunk,
    )
    long_text = "这是一条尽调需求条目示例内容。\n" * 300  # ~30个字符 × 300行 ≈ 9000字符，强制分3块
    from cangjie_fos.services.dd_checklist_parser import _llm_extract_items
    items = _llm_extract_items(long_text)

    assert call_count >= 2, f"长文本应分多块，实际调用次数: {call_count}"
    assert len(items) >= 2


def test_chunked_deduplication_removes_overlap_duplicates(monkeypatch):
    """重叠区域的相同需求项不应出现两次。"""
    def mock_extract_chunk(chunk_text: str) -> list[dict]:
        # 每块都返回同样两条需求（模拟重叠区域重复）
        return [
            {"item_no": "1", "category": "财务", "requirement": "验资报告"},
            {"item_no": "2", "category": "法务", "requirement": "营业执照"},
        ]

    monkeypatch.setattr(
        "cangjie_fos.services.dd_checklist_parser._llm_extract_chunk",
        mock_extract_chunk,
    )
    long_text = "x" * 8000  # 强制分3块
    from cangjie_fos.services.dd_checklist_parser import _llm_extract_items
    items = _llm_extract_items(long_text)

    requirements = [i["requirement"] for i in items]
    assert len(requirements) == len(set(requirements)), "重复需求项应被去重"
    assert len(items) == 2  # "验资报告" 和 "营业执照" 各一条


def test_prefilter_reduces_large_index_to_top_n():
    """超过50个文件时，预筛应只返回最相关的50个（不超限）。"""
    from cangjie_fos.services.dd_match_service import _prefilter_files_for_batch
    index_rows = [
        {"filename": f"文件{i}.pdf", "summary": f"摘要内容{i}"}
        for i in range(100)
    ]
    batch_items = [{"id": "1", "requirement": "财务报告审计年报"}]
    result = _prefilter_files_for_batch(batch_items, index_rows, top_n=50)
    assert len(result) == 50


def test_prefilter_passthrough_when_small_index():
    """文件数不超过 top_n 时，直接返回全量不做筛选。"""
    from cangjie_fos.services.dd_match_service import _prefilter_files_for_batch
    index_rows = [{"filename": f"{i}.pdf", "summary": ""} for i in range(30)]
    batch_items = [{"id": "1", "requirement": "营业执照"}]
    result = _prefilter_files_for_batch(batch_items, index_rows, top_n=50)
    assert len(result) == 30


# ── clean_filename 单元测试 ──────────────────────────────────────

def test_clean_filename_strips_extension():
    from cangjie_fos.services.dd_index_service import clean_filename
    assert clean_filename("验资报告.pdf") == "验资报告"


def test_clean_filename_strips_date_and_version():
    from cangjie_fos.services.dd_index_service import clean_filename
    result = clean_filename("审计报告2024年06月最终版.pdf")
    assert "2024" not in result
    assert "最终版" not in result
    assert "审计报告" in result


def test_clean_filename_strips_version_tag():
    from cangjie_fos.services.dd_index_service import clean_filename
    result = clean_filename("股权结构图v2.1.xlsx")
    assert "v2.1" not in result
    assert "股权结构图" in result


def test_clean_filename_improves_prefilter_match():
    """clean_filename 使带年份的文件名也能匹配到相关需求。"""
    from cangjie_fos.services.dd_match_service import _prefilter_files_for_batch
    # 文件名含年份，不加 clean 时 bigram "审计" 无法命中
    index_rows = [
        {"filename": "2023年度审计报告_盖章版.pdf", "summary": None},
        {"filename": "完全不相关文件.docx", "summary": None},
    ] * 30  # 超过 top_n 阈值
    batch_items = [{"id": "1", "requirement": "审计报告"}]
    result = _prefilter_files_for_batch(batch_items, index_rows, top_n=30)
    filenames = [r["filename"] for r in result]
    assert "2023年度审计报告_盖章版.pdf" in filenames


# ── Top-3 candidates 存储测试 ─────────────────────────────────────

def test_run_matching_stores_candidates_json():
    """匹配结果中应包含 candidates_json，且主文件为置信度最高的候选。"""
    import json
    from cangjie_fos.services.dd_match_service import create_match_session, run_matching, get_session_items

    items = [{"item_no": "1", "category": "财务", "requirement": "审计报告"}]
    session_id = create_match_session("test_tenant", "test.xlsx", "/folder", items)

    fake_index = [
        {"file_path": "/folder/审计.pdf", "filename": "审计.pdf", "summary": "审计报告"},
        {"file_path": "/folder/备选.pdf", "filename": "备选.pdf", "summary": "备选"},
    ]

    def fake_batch_match(items, file_list_text, index_rows):
        return {items[0]["id"]: {
            "file_path": "/folder/审计.pdf",
            "filename": "审计.pdf",
            "confidence": 0.92,
            "reason": "最佳匹配",
            "candidates_json": json.dumps([
                {"file_path": "/folder/审计.pdf", "filename": "审计.pdf", "confidence": 0.92, "reason": "最佳"},
                {"file_path": "/folder/备选.pdf", "filename": "备选.pdf", "confidence": 0.65, "reason": "次选"},
            ], ensure_ascii=False),
        }}

    from unittest.mock import patch
    with patch("cangjie_fos.services.dd_match_service._llm_batch_match", side_effect=fake_batch_match):
        with patch("cangjie_fos.services.dd_match_service._get_index_for_folder", return_value=fake_index):
            run_matching(session_id, "/folder")

    result = get_session_items(session_id)
    assert result[0]["matched_filename"] == "审计.pdf"
    assert result[0]["candidates_json"] is not None
    cands = json.loads(result[0]["candidates_json"])
    assert len(cands) == 2
    assert cands[1]["filename"] == "备选.pdf"
