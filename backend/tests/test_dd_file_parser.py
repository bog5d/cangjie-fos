"""Unit tests for dd_file_parser — no LLM, no network."""
from __future__ import annotations
import io
from pathlib import Path
import pytest


def test_extract_txt_file(tmp_path):
    f = tmp_path / "report.txt"
    f.write_text("这是一份财务报告，包含2023年营收数据。", encoding="utf-8")
    from cangjie_fos.services.dd_file_parser import extract_text
    text, readable = extract_text(f)
    assert readable is True
    assert "财务报告" in text


def test_extract_unsupported_file_returns_not_readable(tmp_path):
    f = tmp_path / "image.png"
    f.write_bytes(b"\x89PNG\r\n")
    from cangjie_fos.services.dd_file_parser import extract_text
    text, readable = extract_text(f)
    assert readable is False
    assert text == ""


def test_extract_excel_file(tmp_path):
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["序号", "内容"])
    ws.append(["1", "验资报告"])
    ws.append(["2", "营业执照"])
    path = tmp_path / "checklist.xlsx"
    wb.save(str(path))

    from cangjie_fos.services.dd_file_parser import extract_text
    text, readable = extract_text(path)
    assert readable is True
    assert "验资报告" in text


def test_extract_docx_file(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_paragraph("公司架构图说明")
    doc.add_paragraph("组织结构详见附件")
    path = tmp_path / "org.docx"
    doc.save(str(path))

    from cangjie_fos.services.dd_file_parser import extract_text
    text, readable = extract_text(path)
    assert readable is True
    assert "公司架构图" in text


# ── dd_index_service 测试 ────────────────────────────────────

def test_scan_indexes_txt_files(tmp_path):
    (tmp_path / "report.txt").write_text("财务审计报告2023年度", encoding="utf-8")
    (tmp_path / "image.png").write_bytes(b"\x89PNG")  # 不支持，应记录 readable=0

    from unittest.mock import patch
    from cangjie_fos.services.dd_index_service import scan_and_index_folder, get_index_by_folder

    with patch("cangjie_fos.services.dd_index_service._llm_summarize", return_value="2023年财务审计报告"):
        result = scan_and_index_folder(str(tmp_path), "test")

    assert result["indexed"] == 1   # txt 成功
    # png 不在 SUPPORTED_EXTENSIONS，不进索引
    rows = get_index_by_folder(str(tmp_path))
    assert len(rows) == 1
    assert rows[0]["filename"] == "report.txt"
    assert rows[0]["summary"] == "2023年财务审计报告"
    assert rows[0]["readable"] == 1


def test_scan_invalid_folder_raises():
    from cangjie_fos.services.dd_index_service import scan_and_index_folder
    with pytest.raises(ValueError, match="Not a directory"):
        scan_and_index_folder("/nonexistent/path/xyz", "test")


# ── Bug 1 回归测试：大文件夹跳过 LLM 摘要 ─────────────────────────────────────

def test_large_folder_skips_llm_summarize(tmp_path):
    """Bug 1: 文件数 > 200 时，_llm_summarize 不应被调用（仅索引文件名）。"""
    from unittest.mock import patch, MagicMock
    from cangjie_fos.services.dd_index_service import (
        scan_and_index_folder,
        get_index_by_folder,
        MAX_LLM_SUMMARIZE_FILES,
    )

    # 创建 MAX_LLM_SUMMARIZE_FILES + 1 个 txt 文件
    n = MAX_LLM_SUMMARIZE_FILES + 1
    for i in range(n):
        (tmp_path / f"file_{i:04d}.txt").write_text(f"内容{i}", encoding="utf-8")

    mock_llm = MagicMock(return_value="mock_summary")
    with patch("cangjie_fos.services.dd_index_service._llm_summarize", mock_llm):
        result = scan_and_index_folder(str(tmp_path), "test_large")

    # LLM 摘要不应该被调用
    mock_llm.assert_not_called()
    assert result["indexed"] == n
    # summary 应全部为 None（未生成摘要）
    rows = get_index_by_folder(str(tmp_path))
    assert all(r["summary"] is None for r in rows)


def test_small_folder_calls_llm_summarize(tmp_path):
    """Bug 1 对照：文件数 <= 200 时，仍应调用 LLM 摘要。"""
    from unittest.mock import patch
    from cangjie_fos.services.dd_index_service import scan_and_index_folder

    (tmp_path / "doc.txt").write_text("审计报告", encoding="utf-8")

    with patch("cangjie_fos.services.dd_index_service._llm_summarize", return_value="审计") as mock_llm:
        scan_and_index_folder(str(tmp_path), "test_small")

    mock_llm.assert_called_once()


def test_scan_progress_callback(tmp_path):
    """Bug 1: progress_callback 应在每 50 个文件后被调用。"""
    from unittest.mock import patch
    from cangjie_fos.services.dd_index_service import scan_and_index_folder

    # 创建 51 个 txt 文件（超过一个 50 文件批次）
    for i in range(51):
        (tmp_path / f"f_{i:03d}.txt").write_text("内容", encoding="utf-8")

    calls: list[tuple[int, int]] = []

    def _progress(done: int, total: int) -> None:
        calls.append((done, total))

    with patch("cangjie_fos.services.dd_index_service._llm_summarize", return_value="摘要"):
        scan_and_index_folder(str(tmp_path), "test_progress", progress_callback=_progress)

    # 51 个文件，第50个触发一次回调
    assert len(calls) >= 1
    # 第一次回调的 total 应为 51
    assert calls[0][1] == 51


# ── Bug 3 回归测试：不可读文件也参与匹配 ─────────────────────────────────────

def test_get_index_for_folder_includes_unreadable_files(tmp_path):
    """Bug 3: _get_index_for_folder 不应过滤 readable=0 的文件。"""
    from unittest.mock import patch
    from cangjie_fos.services.dd_index_service import scan_and_index_folder
    from cangjie_fos.services.dd_match_service import _get_index_for_folder

    # txt（readable=1）+ 图片型PDF（readable=0，无法提取文字）
    (tmp_path / "report.txt").write_text("财务报告内容", encoding="utf-8")

    # 模拟一个 readable=False 的文件（extract_text 返回 readable=False）
    with patch("cangjie_fos.services.dd_index_service._llm_summarize", return_value="财务"):
        with patch("cangjie_fos.services.dd_index_service.extract_text") as mock_ext:
            # 第一个文件 readable=True，第二个 readable=False
            mock_ext.side_effect = [
                ("财务报告内容", True),
                ("", False),
            ]
            # 手动调用 _index_single_file 来构造两条索引
            from cangjie_fos.services.dd_index_service import _index_single_file
            _index_single_file(tmp_path / "report.txt", str(tmp_path), use_llm=True)

            # 插入一条 readable=0 的记录（模拟图片PDF）
            import time
            import uuid
            from cangjie_fos.services.db_base import _connect
            with _connect() as conn:
                conn.execute(
                    """INSERT INTO dd_asset_index
                       (id, folder_root, file_path, filename, file_type, summary, readable, indexed_at)
                       VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
                    (str(uuid.uuid4()), str(tmp_path),
                     str(tmp_path / "scan_img.pdf"), "scan_img.pdf",
                     ".pdf", None, 0, time.time()),
                )

    rows = _get_index_for_folder(str(tmp_path))
    filenames = [r["filename"] for r in rows]

    # 两个文件都应该在结果里（不过滤 readable=0）
    assert "report.txt" in filenames
    assert "scan_img.pdf" in filenames
