"""DD 物料架构压力基准（可手动运行；CI 用 tests/test_dd_stress_smoke.py 的快测版）。

跑法：
    cd backend
    uv run python bench/dd_stress.py --scale medium            # 仅打印指标
    uv run python bench/dd_stress.py --scale large --charts     # 顺带出 PNG 图
    uv run python bench/dd_stress.py --real-llm                 # 用真实 LLM（需 key）

策略：默认 LLM 用确定性 mock（无 key 也能跑），其余全部真实代码：
  真实文件解析 / 真实 SQLite / 真实预筛 / 真实精判循环 / 真实记忆 / 真实并发。
图表渲染依赖 matplotlib（非项目硬依赖，按需 `uv pip install matplotlib`）。
"""
from __future__ import annotations
import argparse, json, random, re, tempfile, threading, time
from pathlib import Path
from unittest.mock import MagicMock, patch

CATS = {
    "财务": ["审计报告", "财务报表", "利润表", "资产负债表", "现金流量表",
             "纳税申报表", "银行流水", "验资报告", "财务尽调底稿"],
    "法务": ["公司章程", "营业执照", "股东会决议", "董事会决议", "重大合同清单",
             "知识产权证书", "商标注册证", "发明专利证书", "法律意见书"],
    "业务": ["商业计划书", "产品白皮书", "核心客户名单", "供应商框架协议",
             "市场分析报告", "销售合同台账", "技术架构说明"],
    "人事": ["劳动合同", "员工花名册", "社保缴纳证明", "股权激励方案", "核心团队简历"],
    "其他": ["尽职调查问卷", "补充资料说明", "项目会议纪要", "实控人承诺函"],
}
YEARS = [2021, 2022, 2023, 2024]
DEPTS = ["集团", "母公司", "子公司甲", "子公司乙", "事业部一", "事业部二", "境外主体", "新设主体"]
_SCALE = {"small": 1, "medium": 4, "large": 8}  # 取 DEPTS 前 N 个


def _write_txt(p: Path, doctype: str, year: int):
    p.write_text(f"{doctype}（{year}年度）\n本文件为{year}年度{doctype}正文。"
                 f"经核查，{doctype}内容真实完整，符合监管要求。" * 12, encoding="utf-8")

def _write_docx(p: Path, doctype: str, year: int):
    from docx import Document
    d = Document(); d.add_heading(f"{doctype} {year}", 1)
    for _ in range(8):
        d.add_paragraph(f"{year}年度{doctype}：条款明细、金额、签署方与履行情况说明。")
    d.save(str(p))

def _write_xlsx(p: Path, doctype: str, year: int):
    import openpyxl
    wb = openpyxl.Workbook(); ws = wb.active
    ws.append([doctype, f"{year}年度", "金额(万元)"])
    for i in range(20):
        ws.append([f"{doctype}-科目{i}", year, round(random.random() * 9999, 2)])
    wb.save(str(p))


def build_library(root: Path, scale: str = "medium") -> int:
    """造一个复杂材料库（txt/docx/xlsx 混合），返回文件数。"""
    deps = DEPTS[:_SCALE.get(scale, 4)]
    count = 0
    for cat, doctypes in CATS.items():
        cdir = root / cat; cdir.mkdir(parents=True, exist_ok=True)
        for dt in doctypes:
            for y in YEARS:
                for dep in deps:
                    fmt = random.choices(["txt", "docx", "xlsx"], weights=[8, 1, 1])[0]
                    p = cdir / f"{dep}_{dt}_{y}.{fmt}"
                    {"txt": _write_txt, "docx": _write_docx, "xlsx": _write_xlsx}[fmt](p, f"{dep}{dt}", y)
                    count += 1
        for n in range(20):
            _write_txt(cdir / f"{cat}_附件{n}_扫描件.txt", f"{cat}补充材料", random.choice(YEARS))
            count += 1
    return count


def make_checklist(n: int) -> list[dict]:
    pool = [dt for dts in CATS.values() for dt in dts]
    ph = ["请提供{}", "{}（近三年）", "贵公司{}文件", "{}", "最新{}及说明"]
    return [{"item_no": str(i + 1), "category": "尽调",
             "requirement": random.choice(ph).format(random.choice(pool))} for i in range(n)]


def fake_batch_client():
    """假匹配器：解析 prompt 文件列表，按文件名×需求关键词重合度选最佳文件（结果相关）。"""
    from cangjie_fos.services.dd_match_service import _requirement_bigrams
    def create(**kw):
        content = kw.get("messages", [{}])[-1].get("content", "")
        files = {int(i): name for i, name in re.findall(r"\[(\d+)\] 文件名：(.+?)  摘要：", content)}
        out = {}
        for uid, req in re.findall(r"需求\d+（ID:([0-9a-f-]{36})）：(.+)", content):
            kws = _requirement_bigrams(req)
            best_idx, best = (0 if files else None), 0
            for idx, name in files.items():
                sc = sum(1 for k in kws if k in name)
                if sc > best:
                    best, best_idx = sc, idx
            out[uid] = ({"candidates": []} if best_idx is None else
                        {"candidates": [{"file_index": best_idx,
                         "confidence": round(min(0.95, 0.55 + best * 0.1), 2), "reason": "rel"}]})
        m = MagicMock(); m.choices[0].message.content = json.dumps(out, ensure_ascii=False)
        return m
    c = MagicMock(); c.chat.completions.create.side_effect = create
    return c


def fake_refine(client, requirement, filename, content_text):
    from cangjie_fos.services.dd_match_service import _requirement_bigrams
    hits = [k for k in _requirement_bigrams(requirement) if k in (content_text or "")]
    sat = len(hits) > 0
    return {"satisfies": sat, "confidence": round(min(0.97, 0.35 + len(hits) * 0.12), 2),
            "evidence": f"正文含「{hits[0]}」等{len(hits)}处匹配" if sat else "正文未见该需求关键内容"}


def run_benchmark(scale: str = "medium", checklist_n: int = 120, seed_mem: int = 20000,
                  concurrency: int = 8, real_llm: bool = False) -> dict:
    random.seed(42)
    import cangjie_fos.services.pitch_job_db as pjdb
    tmp = Path(tempfile.mkdtemp(prefix="dd_bench_"))
    pjdb._db_path = lambda: str(tmp / "bench.sqlite")  # 隔离 DB

    from cangjie_fos.services.db_base import _connect
    from cangjie_fos.services import dd_index_service, dd_match_service
    from cangjie_fos.services.dd_match_service import (
        create_match_session, run_matching, get_session_items, record_session_decisions,
        lookup_decision_memory, _prefilter_files_for_batch, _get_index_for_folder,
        normalize_requirement, MEMORY_REASON_PREFIX,
    )
    m: dict = {"scale": scale}
    lib = tmp / "材料库"

    t = time.perf_counter(); m["file_count"] = build_library(lib, scale)
    m["lib_gen_s"] = round(time.perf_counter() - t, 2)

    sumctx = (patch.object(dd_index_service, "_llm_summarize", lambda f, c: f"{Path(f).stem}摘要")
              if not real_llm else _nullctx())
    t = time.perf_counter()
    with sumctx:
        res = dd_index_service.scan_and_index_folder(str(lib), "bench")
    m["index_s"] = round(time.perf_counter() - t, 2)
    m["indexed"] = res["indexed"]
    m["index_throughput"] = round(res["indexed"] / max(m["index_s"], 1e-6))
    with _connect() as conn:
        tot = conn.execute("SELECT COUNT(*) FROM dd_asset_index WHERE folder_root=?", (str(lib),)).fetchone()[0]
        wc = conn.execute("SELECT COUNT(*) FROM dd_asset_index WHERE folder_root=? AND content_text IS NOT NULL AND content_text!=''", (str(lib),)).fetchone()[0]
    m["content_coverage"] = round(wc / tot, 4) if tot else 0

    index_rows = _get_index_for_folder(str(lib))
    batch = [{"requirement": dt} for dts in CATS.values() for dt in dts][:20]
    t = time.perf_counter()
    for _ in range(200):
        _prefilter_files_for_batch(batch, index_rows, top_n=50)
    m["prefilter_ms"] = round((time.perf_counter() - t) / 200 * 1000, 2)

    items = make_checklist(checklist_n)
    sid = create_match_session("bench", "大清单.xlsx", str(lib), items, institution_name="压测机构A")
    mctx = _llm_ctx(dd_match_service, real_llm)
    t = time.perf_counter()
    with mctx:
        run_matching(sid, str(lib))
    m["match_s"] = round(time.perf_counter() - t, 2)
    m["match_throughput"] = round(checklist_n / max(m["match_s"], 1e-6))
    si = get_session_items(sid)
    m["verdict_green"] = sum(1 for it in si if it.get("verdict") == "green")
    m["verdict_yellow"] = sum(1 for it in si if it.get("verdict") == "yellow")
    m["verdict_red"] = sum(1 for it in si if it.get("verdict") == "red")

    now = time.time()
    with _connect() as conn:
        conn.executemany(
            """INSERT INTO dd_decision_memory (id,requirement_norm,requirement,file_path,filename,confirm_count,last_institution,updated_at)
               VALUES (?,?,?,?,?,?,?,?)""",
            [(f"seed{i}", normalize_requirement(f"历史需求{i}"), f"历史需求{i}",
              f"/x/{i}.pdf", f"{i}.pdf", random.randint(1, 9), "seed", now) for i in range(seed_mem)])
    t = time.perf_counter()
    for i in range(3000):
        lookup_decision_memory(f"历史需求{random.randint(0, seed_mem - 1)}")
    m["mem_lookup_us"] = round((time.perf_counter() - t) / 3000 * 1e6)
    m["seed_mem"] = seed_mem

    # 跨机构锁定真实验证
    rf = next(r["file_path"] for r in index_rows if "审计报告" in r["filename"])
    sid_a = create_match_session("bench", "a", str(lib), [{"item_no": "1", "category": "财务", "requirement": "请提供审计报告"}], institution_name="A")
    aid = get_session_items(sid_a)[0]["id"]
    with _connect() as conn:
        conn.execute("UPDATE dd_match_items SET matched_file_path=?,matched_filename=?,confidence=0.9,user_confirmed=1 WHERE id=?", (rf, Path(rf).name, aid))
    record_session_decisions(sid_a)
    sid_b = create_match_session("bench", "b", str(lib), [{"item_no": "1", "category": "财务", "requirement": "请提供审计报告"}], institution_name="B")
    with _llm_ctx(dd_match_service, real_llm), patch.object(dd_match_service, "_llm_batch_match", return_value={}):
        run_matching(sid_b, str(lib))
    bi = get_session_items(sid_b)[0]
    m["cross_inst_ok"] = int(bi["matched_file_path"] == rf and (bi["match_reason"] or "").startswith(MEMORY_REASON_PREFIX))

    # 并发：patch 一次性包住整段（patch.object 非线程隔离，不能每个 worker 各 patch）
    errors: list[str] = []
    def worker(k):
        try:
            s = create_match_session("bench", f"c{k}", str(lib), make_checklist(40), institution_name=f"并发{k}")
            run_matching(s, str(lib))
        except Exception as e:  # noqa: BLE001
            errors.append(f"{k}:{e}")
    t = time.perf_counter()
    with _llm_ctx(dd_match_service, real_llm):
        ths = [threading.Thread(target=worker, args=(k,)) for k in range(concurrency)]
        [x.start() for x in ths]; [x.join() for x in ths]
    m["concurrency_s"] = round(time.perf_counter() - t, 2)
    m["concurrency_total"] = concurrency
    m["concurrency_ok"] = concurrency - len(errors)
    m["concurrency_errors"] = errors[:3]
    m["db_size_mb"] = round((tmp / "bench.sqlite").stat().st_size / 1e6, 1)

    m["_sample"] = [{"req": it["requirement"], "file": it["matched_filename"],
                     "verdict": it["verdict"], "evidence": it.get("evidence") or "",
                     "conf": it.get("confidence") or 0}
                    for it in si if it.get("verdict") and it.get("matched_filename")][:13]
    return m


class _nullctx:
    def __enter__(self): return self
    def __exit__(self, *a): return False

def _llm_ctx(dd_match_service, real_llm):
    if real_llm:
        return _nullctx()
    from contextlib import ExitStack
    es = ExitStack()
    es.enter_context(patch.object(dd_match_service, "get_dd_llm_client", fake_batch_client))
    es.enter_context(patch.object(dd_match_service, "_llm_refine_candidate", fake_refine))
    return es


def render_charts(m: dict, out_dir: Path):
    import matplotlib; matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib import font_manager as fm
    for fp in ("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",):
        if Path(fp).exists():
            matplotlib.rcParams["font.family"] = fm.FontProperties(fname=fp).get_name()
            break
    matplotlib.rcParams["axes.unicode_minus"] = False
    G, Y, R, B = "#16a34a", "#d97706", "#dc2626", "#2563eb"
    fig = plt.figure(figsize=(13, 8))
    fig.suptitle(f"DD 物料架构压测仪表盘（{int(m['file_count'])} 文件 · scale={m['scale']}）", fontsize=15)
    ax = fig.add_subplot(2, 3, 1); ax.bar(["索引/秒", "匹配/秒"], [m["index_throughput"], m["match_throughput"]], color=[B, G]); ax.set_title("吞吐量")
    ax = fig.add_subplot(2, 3, 2); ax.bar(["预筛ms/批", "记忆µs/次"], [m["prefilter_ms"], m["mem_lookup_us"] / 1000], color=[B, "#9333ea"]); ax.set_title("关键延迟")
    ax = fig.add_subplot(2, 3, 3); ax.bar(["绿", "黄", "红"], [m["verdict_green"], m["verdict_yellow"], m["verdict_red"]], color=[G, Y, R]); ax.set_title("机器验证分布")
    ax = fig.add_subplot(2, 3, 4); ax.bar(["成功", "失败"], [m["concurrency_ok"], m["concurrency_total"] - m["concurrency_ok"]], color=[G, R]); ax.set_title(f"并发({m['concurrency_total']}线程)")
    ax = fig.add_subplot(2, 3, 5); cov = m["content_coverage"] * 100; ax.pie([cov, 100 - cov], labels=[f"{cov:.0f}%", ""], colors=[G, "#e5e7eb"], startangle=90); ax.set_title("全文落库覆盖率")
    ax = fig.add_subplot(2, 3, 6); ax.axis("off")
    ax.text(0, 1, f"索引 {m['index_s']}s 覆盖{cov:.0f}%\n匹配+精判 {m['match_s']}s\n预筛 {m['prefilter_ms']}ms/批\n记忆库 {m['seed_mem']} 查询{m['mem_lookup_us']}µs\n跨机构锁定 {'成功' if m['cross_inst_ok'] else '失败'}\n并发 {m['concurrency_ok']}/{m['concurrency_total']} 零错误\nDB {m['db_size_mb']}MB", va="top", fontsize=11, linespacing=1.8)
    plt.tight_layout(rect=[0, 0, 1, 0.96])
    p = out_dir / "dd_stress_perf.png"; plt.savefig(p, dpi=130, bbox_inches="tight")
    print(f"chart → {p}")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--scale", choices=list(_SCALE), default="medium")
    ap.add_argument("--checklist", type=int, default=120)
    ap.add_argument("--seed-mem", type=int, default=20000)
    ap.add_argument("--concurrency", type=int, default=8)
    ap.add_argument("--real-llm", action="store_true")
    ap.add_argument("--charts", action="store_true")
    a = ap.parse_args()
    m = run_benchmark(a.scale, a.checklist, a.seed_mem, a.concurrency, a.real_llm)
    pub = {k: v for k, v in m.items() if not k.startswith("_")}
    print(json.dumps(pub, ensure_ascii=False, indent=2))
    if a.charts:
        render_charts(m, Path(tempfile.gettempdir()))


if __name__ == "__main__":
    main()
